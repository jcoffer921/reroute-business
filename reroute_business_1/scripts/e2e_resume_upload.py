"""
End-to-end sanity test for the resume upload endpoint using Django's test client.
Creates a user, logs in, uploads a generated DOCX, and verifies JSON response
and imported resume page load.
"""

import io
import json
import os
import sys

# Ensure project root (parent of scripts/) is on sys.path
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
if BASE_DIR not in sys.path:
    sys.path.insert(0, BASE_DIR)

# Force DEBUG=true to avoid manifest static lookups and HTTPS redirects
os.environ.setdefault("DEBUG", "true")
os.environ.setdefault("DJANGO_SETTINGS_MODULE", "reroute.settings")

import django  # noqa: E402
django.setup()
from django.conf import settings  # noqa: E402
if 'testserver' not in settings.ALLOWED_HOSTS:
    settings.ALLOWED_HOSTS = list(settings.ALLOWED_HOSTS) + ['testserver']

from django.contrib.auth.models import User  # noqa: E402
from django.core.files.uploadedfile import SimpleUploadedFile  # noqa: E402
from django.test import Client  # noqa: E402


def main() -> int:
    # Ensure test user exists
    username = "e2e_uploader"
    password = "pass1234!"
    user, created = User.objects.get_or_create(username=username, defaults={
        "email": "e2e@example.com",
    })
    if created:
        user.set_password(password)
        user.save()

    # Disable HTTPS redirect in tests
    settings.SECURE_SSL_REDIRECT = False
    client = Client()
    assert client.login(username=username, password=password), "Login failed"

    # Build a tiny DOCX in-memory
    try:
        from docx import Document  # type: ignore
    except Exception as e:
        print("python-docx not installed or import failed:", e)
        return 2

    bio = io.BytesIO()
    doc = Document()
    doc.add_paragraph("John Test")
    doc.add_paragraph("Email: john@example.com")
    doc.add_paragraph("Skills: forklift, punctual, safety")
    doc.save(bio)
    bio.seek(0)

    upload = SimpleUploadedFile(
        "test.docx",
        bio.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    # Hit the upload endpoint
    resp = client.post("/resume/parse-upload/", data={"file": upload})
    print("POST /resume/parse-upload/ ->", resp.status_code)
    try:
        payload = json.loads(resp.content.decode("utf-8"))
    except Exception:
        print("Non-JSON response:", resp.content[:500])
        return 3

    print("Response JSON:", payload)
    if resp.status_code != 200 or "resume_id" not in payload:
        return 4

    # Fetch the imported resume page
    resume_id = payload["resume_id"]
    resp2 = client.get(f"/resume/import/{resume_id}/")
    print(f"GET /resume/import/{resume_id}/ ->", resp2.status_code)
    if resp2.status_code != 200:
        print(resp2.content[:500])
        return 5

    print("E2E upload flow looks good.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
